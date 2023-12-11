from docx import Document
from docx.enum.text import WD_COLOR_INDEX

doc = Document('Sample1.docx')

# Get the first paragraph's text
p1_text = doc.paragraphs[0].text

# Create a new paragraph with "excellent" highlighted
p2 = doc.add_paragraph()
substrings = p1_text.split('excellent')
for substring in substrings[:-1]:
   # p1_text.add_run(substring)
    font = p2.add_run('excellent').font
    font.highlight_color = WD_COLOR_INDEX.YELLOW
p2.add_run(substrings[-1])

# Save document under new name
doc.save('t1.docx')