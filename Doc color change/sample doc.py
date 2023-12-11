from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
doc=Document()

run = doc.add_paragraph().add_run('A random paragraph can also be an excellent way for a writer to tackle writers block. Writing block can often happen due to being stuck with a current project that the writer is trying to complete. By inserting a completely random paragraph from which to begin, it can take down some of the issues that may have been causing the writers block in the first place.')
font = run.font
font.color.rgb = RGBColor(0x20, 0x80, 0x20)


doc.add_paragraph().add_run('A random paragraph can also be an excellent way for a writer to tackle writers block. Writing block can often happen due to being stuck with a current project that the writer is trying to complete. By inserting a completely random paragraph from which to begin, it can take down some of the issues that may have been causing the writers block in the first place.'
                  ).font.highlight_color = WD_COLOR_INDEX.BLUE

doc.save('D:\Web Scraping Tute\Doc color change\Sample.docx')
